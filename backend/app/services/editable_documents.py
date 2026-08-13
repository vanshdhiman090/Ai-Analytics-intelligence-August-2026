"""Structured, versionable document drafts with editable Word and PDF exports."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Sequence
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360


def _text(value: Any, fallback: str = "Not available") -> str:
    if value is None or value == "":
        return fallback
    return str(value)


def _items(values: Iterable[Any] | None) -> list[str]:
    return [str(value) for value in (values or []) if value not in (None, "")]


def _table(title: str, columns: Sequence[str], rows: Iterable[Sequence[Any]]) -> dict:
    return {
        "type": "table",
        "title": title,
        "columns": list(columns),
        "rows": [[_text(value, "—") for value in row] for row in rows],
    }


def _bullets(title: str, values: Iterable[Any] | None, fallback: str = "None recorded") -> dict:
    items = _items(values)
    return {"type": "bullets", "title": title, "items": items or [fallback]}


def _prose(title: str, value: Any, fallback: str = "Not recorded") -> dict:
    return {"type": "prose", "title": title, "text": _text(value, fallback)}


def _profile_columns(profile: dict) -> list[dict]:
    columns = profile.get("columns") or {}
    if isinstance(columns, dict):
        return [dict(value, name=value.get("name", name)) for name, value in columns.items()]
    return list(columns)


def seed_document(title: str, session: Any, dataset: Any, checkpoints: list[Any], actions: list[Any], document_type_hint: str | None = None) -> dict:
    """Build a rich version-zero editor document from the validated run package."""
    result = session.result_summary or {}
    profile = result.get("schema_profile") or ((dataset.schema_profile if dataset is not None else None) or {})
    brief = result.get("analysis_brief") or {}
    lower_title = title.lower()
    document_type = document_type_hint or (
        "analysis_journal" if "journal" in lower_title
        else "technical_appendix" if "technical" in lower_title or "appendix" in lower_title
        else "analysis_brief" if "brief" in lower_title
        else "executive_report"
    )

    findings = result.get("findings") or []
    recommendations = result.get("recommendations") or []
    evidence = result.get("evidence") or []
    source_register = result.get("source_register") or []
    cleaning_log = result.get("cleaning_log") or []
    integrity_checks = result.get("integrity_checks") or []
    quality_gates = result.get("quality_gates") or []
    operations = (result.get("analysis_plan") or {}).get("operations") or []

    finding_table = _table("Question → finding → implication → evidence", ["ID", "Finding", "Implication", "Evidence", "Confidence"], (
        [item.get("finding_id"), item.get("statement"), item.get("implication"), ", ".join(item.get("evidence_ids") or []), item.get("confidence", "Unknown")]
        for item in findings
    ))
    recommendation_table = _table("Evidence-linked action plan", ["ID", "Action", "Rationale", "Owner", "Timeframe", "Impact", "Effort", "Findings"], (
        [item.get("recommendation_id"), item.get("action"), item.get("rationale"), item.get("owner_role"), item.get("timeframe"), item.get("expected_impact", "Unknown"), item.get("effort", "Unknown"), ", ".join(item.get("finding_ids") or [])]
        for item in recommendations
    ))
    gate_table = _table("Publication quality gates", ["ID", "Gate", "Status", "Severity", "Result"], (
        [item.get("gate_id"), item.get("name"), item.get("status"), item.get("severity"), item.get("detail")]
        for item in quality_gates
    ))
    measurement_table = _table("Measurement scope", ["Evidence", "Measure", "Method", "Population"], (
        [item.get("evidence_id"), item.get("title"), item.get("method"), item.get("population")]
        for item in evidence
    ))
    first_action = recommendations[0].get("action") if recommendations else "Monitor the stated metrics and resolve open questions before acting."

    if document_type == "case_study_report":
        sections = [
            {"heading": "Executive Summary", "body": "", "blocks": [
                _bullets("Decision-ready summary", [
                    f"Bottom line: {result.get('summary') or 'No executive summary was produced.'}",
                    f"Decision supported: {_text(brief.get('decision'), 'The decision owner was not specified.')}",
                    f"Recommended response: {first_action}",
                ])
            ]},
            {"phase": "Ask", "heading": "01 Â· Ask â€” Business task and success criteria", "body": "", "blocks": [
                _prose("Clear business task", brief.get("primary_question", session.business_task)),
                _prose("Objective", brief.get("objective")),
                _prose("Decision supported", brief.get("decision")),
                _table("Stakeholders and audience", ["Name", "Role", "Decision interest"], ([item.get("name"), item.get("role"), item.get("decision_interest")] for item in brief.get("stakeholders") or [])),
                _bullets("Success criteria", brief.get("success_criteria")),
                _bullets("In scope", brief.get("in_scope")),
                _bullets("Out of scope", brief.get("out_of_scope")),
            ]},
            {"phase": "Prepare", "heading": "02 Â· Prepare â€” Data sources and credibility", "body": "", "blocks": [
                _table("All data sources used", ["ID", "File", "Format", "Rows", "Columns", "Observed grain", "Source / licence"], ([item.get("source_id"), item.get("filename"), item.get("format"), item.get("rows"), item.get("columns"), item.get("grain"), item.get("licence")] for item in source_register)),
                _table("ROCCC assessment", ["Dimension", "Recorded assessment"], ([key.replace("_", " ").title(), value] for key, value in (result.get("roccc_answers") or {}).items())),
                _table("Field inventory", ["Field", "Stored type", "Semantic type", "Nulls", "Null %", "Unique", "Examples"], ([item.get("name"), item.get("dtype"), item.get("semantic_type"), item.get("null_count"), item.get("null_pct"), item.get("unique_count"), ", ".join(map(str, item.get("sample_values") or []))] for item in _profile_columns(profile))),
            ]},
            {"phase": "Process", "heading": "03 Â· Process â€” Cleaning and manipulation record", "body": "", "blocks": [
                _prose("Validation status", result.get("validation_status", "Unknown")),
                _table("Cleaning log", ["ID", "Action", "Columns", "Affected", "Before", "After", "Reason"], ([item.get("log_id"), item.get("action"), item.get("columns"), item.get("rows_affected"), item.get("before"), item.get("after"), item.get("reason")] for item in cleaning_log)),
                _table("Integrity checks", ["ID", "Check", "Status", "Result", "Analytical implication"], ([item.get("check_id"), item.get("check"), item.get("status"), item.get("detail"), item.get("implication")] for item in integrity_checks)),
                _bullets("Quality findings", [f"{item.get('column', 'Dataset')}: {item.get('issue')} ({item.get('severity', 'unknown')})" for item in result.get("quality_findings") or []]),
            ]},
            {"phase": "Analyze", "heading": "04 Â· Analyze â€” Summary, calculations, and findings", "body": "", "blocks": [
                _prose("Summary of the analysis", result.get("summary")),
                _table("Approved calculation plan", ["ID", "Operation", "Metric", "Dimension", "Time", "Aggregate", "Rationale"], ([item.get("operation_id"), item.get("kind"), item.get("metric_column"), item.get("dimension_column"), item.get("time_column"), item.get("aggregation"), item.get("rationale")] for item in operations)),
                finding_table,
                _table("Evidence register", ["ID", "Result", "Operation", "Quality", "Method", "Population", "Caveats"], ([item.get("evidence_id"), item.get("title"), item.get("kind"), item.get("quality_status", "ready"), item.get("method"), item.get("population"), "; ".join(item.get("caveats") or [])] for item in evidence)),
            ]},
            {"phase": "Share", "heading": "05 Â· Share â€” Visualizations and key findings", "body": "The validated evidence charts are embedded after the editable narrative sections.", "blocks": [finding_table]},
            {"phase": "Act", "heading": "06 Â· Act â€” High-level insights and action", "body": "", "blocks": [
                _bullets("Top high-level insights", [item.get("statement") for item in findings]),
                recommendation_table,
                _bullets("Monitoring plan", result.get("monitoring_metrics")),
                _bullets("Further exploration", result.get("unanswered_questions")),
                _bullets("Caveats and assumptions", result.get("limitations")),
            ]},
            {"heading": "Release readiness", "body": "", "blocks": [gate_table]},
        ]
        subtitle = "One complete, editable Google-style case study following Askâ€“Prepareâ€“Processâ€“Analyzeâ€“Shareâ€“Act"
    elif document_type == "analysis_journal":
        sections = [
            {"phase": "Ask", "heading": "01 · Ask — Business framing", "body": "", "blocks": [
                _prose("Primary question", session.business_task),
                _prose("Objective", brief.get("objective")),
                _prose("Decision supported", brief.get("decision")),
                _table("Stakeholders", ["Name", "Role", "Decision interest"], ([item.get("name"), item.get("role"), item.get("decision_interest")] for item in brief.get("stakeholders") or [])),
                _bullets("Success criteria", brief.get("success_criteria")),
                _bullets("In scope", brief.get("in_scope")),
                _bullets("Out of scope", brief.get("out_of_scope")),
                _bullets("Required human context", brief.get("required_context")),
            ]},
            {"phase": "Prepare", "heading": "02 · Prepare — Source and readiness", "body": "", "blocks": [
                _table("Source register", ["ID", "File", "Format", "Rows", "Columns", "Observed grain", "Candidate keys", "Source / licence"], ([item.get("source_id"), item.get("filename"), item.get("format"), item.get("rows"), item.get("columns"), item.get("grain"), ", ".join(item.get("candidate_keys") or []), item.get("licence")] for item in source_register)),
                _table("ROCCC assessment", ["Dimension", "Recorded assessment"], ([key.replace("_", " ").title(), value] for key, value in (result.get("roccc_answers") or {}).items())),
                _table("Data inventory", ["Field", "Stored type", "Semantic type", "Nulls", "Null %", "Unique", "Example values"], ([item.get("name"), item.get("dtype"), item.get("semantic_type"), item.get("null_count"), item.get("null_pct"), item.get("unique_count"), ", ".join(map(str, item.get("sample_values") or []))] for item in _profile_columns(profile))),
            ]},
            {"phase": "Process", "heading": "03 · Process — Cleaning and validation", "body": "", "blocks": [
                _prose("Validation status", result.get("validation_status", "Unknown")),
                _table("Cleaning log", ["ID", "Action", "Columns", "Rows / cells affected", "Before", "After", "Reason"], ([item.get("log_id"), item.get("action"), item.get("columns"), item.get("rows_affected"), item.get("before"), item.get("after"), item.get("reason")] for item in cleaning_log)),
                _table("Integrity checks", ["ID", "Check", "Status", "Result", "Analytical implication"], ([item.get("check_id"), item.get("check"), item.get("status"), item.get("detail"), item.get("implication")] for item in integrity_checks)),
                _bullets("Quality findings", [f"{item.get('column', 'Dataset')}: {item.get('issue')} ({item.get('severity', 'unknown')})" for item in result.get("quality_findings") or []]),
            ]},
            {"phase": "Analyze", "heading": "04 · Analyze — Calculations and findings", "body": "", "blocks": [
                _table("Approved calculation plan", ["ID", "Operation", "Metric", "Dimension", "Time", "Aggregate", "Rationale"], ([item.get("operation_id"), item.get("kind"), item.get("metric_column"), item.get("dimension_column"), item.get("time_column"), item.get("aggregation"), item.get("rationale")] for item in operations)),
                finding_table,
                _table("Evidence register", ["ID", "Result", "Operation", "Quality", "Method", "Population", "Caveats"], ([item.get("evidence_id"), item.get("title"), item.get("kind"), item.get("quality_status", "ready"), item.get("method"), item.get("population"), "; ".join(item.get("caveats") or [])] for item in evidence)),
            ]},
            {"phase": "Share", "heading": "05 · Share — Communication record", "body": "", "blocks": [
                _prose("Executive narrative", result.get("summary")),
                _bullets("Visual evidence", [action.output_summary or action.action_type for action in actions if action.stage == "share"], "No chart was suitable for the executed operations."),
                _prose("Accessibility note", "Charts retain evidence IDs, descriptive captions, and editable document context. Verify colours and alternative text before external distribution."),
            ]},
            {"phase": "Act", "heading": "06 · Act — Decision follow-through", "body": "", "blocks": [
                recommendation_table,
                _bullets("Monitoring metrics", result.get("monitoring_metrics")),
                _bullets("Limitations", result.get("limitations")),
                _bullets("Open questions", result.get("unanswered_questions")),
            ]},
            {"phase": "Package", "heading": "Neutral package validation", "body": "", "blocks": [gate_table]},
        ]
        subtitle = "Editable phase-by-phase audit journal with traceability, cleaning evidence, integrity checks, and reflection prompts"
    elif document_type == "technical_appendix":
        sections = [
            {"heading": "Source register and provenance", "body": "", "blocks": [_table("Source register", ["ID", "File", "Format", "Rows", "Columns", "Grain", "Candidate keys", "Source / licence"], ([item.get("source_id"), item.get("filename"), item.get("format"), item.get("rows"), item.get("columns"), item.get("grain"), ", ".join(item.get("candidate_keys") or []), item.get("licence")] for item in source_register))]},
            {"heading": "Field dictionary", "body": "", "blocks": [_table("Profiled fields", ["Field", "Type", "Semantic type", "Nulls", "Null %", "Unique", "Min", "Max", "Examples"], ([item.get("name"), item.get("dtype"), item.get("semantic_type"), item.get("null_count"), item.get("null_pct"), item.get("unique_count"), item.get("min"), item.get("max"), ", ".join(map(str, item.get("sample_values") or []))] for item in _profile_columns(profile)))]},
            {"heading": "Transformation and integrity record", "body": "", "blocks": [
                _table("Cleaning log", ["ID", "Action", "Columns", "Affected", "Before", "After", "Reason"], ([item.get("log_id"), item.get("action"), item.get("columns"), item.get("rows_affected"), item.get("before"), item.get("after"), item.get("reason")] for item in cleaning_log)),
                _table("Integrity checks", ["ID", "Check", "Status", "Result", "Implication"], ([item.get("check_id"), item.get("check"), item.get("status"), item.get("detail"), item.get("implication")] for item in integrity_checks)),
            ]},
            {"heading": "Controlled methods", "body": "", "blocks": [_table("Approved analysis operations", ["ID", "Kind", "Metric", "Dimension", "Time", "Aggregate", "Rationale"], ([item.get("operation_id"), item.get("kind"), item.get("metric_column"), item.get("dimension_column"), item.get("time_column"), item.get("aggregation"), item.get("rationale")] for item in operations))]},
            {"heading": "Evidence and traceability", "body": "", "blocks": [finding_table, _table("Evidence register", ["ID", "Title", "Kind", "Quality", "Method", "Population", "Caveats"], ([item.get("evidence_id"), item.get("title"), item.get("kind"), item.get("quality_status", "ready"), item.get("method"), item.get("population"), "; ".join(item.get("caveats") or [])] for item in evidence))]},
            {"heading": "Publication controls", "body": "", "blocks": [gate_table, _bullets("Limitations", result.get("limitations")), _bullets("Unanswered questions", result.get("unanswered_questions"))]},
        ]
        subtitle = "Editable reproducibility companion: provenance, dictionary, cleaning log, integrity tests, methods, evidence, and quality gates"
    elif document_type == "analysis_brief":
        sections = [
            {"heading": "Business purpose", "body": "", "blocks": [_prose("Objective", brief.get("objective", session.business_task)), _prose("Decision supported", brief.get("decision")), _prose("Primary analytical question", brief.get("primary_question", session.business_task))]},
            {"heading": "Stakeholders and success", "body": "", "blocks": [_table("Stakeholders", ["Name", "Role", "Decision interest"], ([item.get("name"), item.get("role"), item.get("decision_interest")] for item in brief.get("stakeholders") or [])), _bullets("Success criteria", brief.get("success_criteria"))]},
            {"heading": "Scope and boundaries", "body": "", "blocks": [_bullets("In scope", brief.get("in_scope")), _bullets("Out of scope", brief.get("out_of_scope")), _bullets("Assumptions", brief.get("assumptions")), _bullets("Constraints", brief.get("constraints")), _bullets("Required context", brief.get("required_context"))]},
        ]
        subtitle = "Editable business contract: decision context, stakeholders, scope, boundaries, and success criteria"
    else:
        sections = [
            {"heading": "Executive Summary", "body": "", "blocks": [
                _bullets("Decision-ready summary", [
                    f"Bottom line: {result.get('summary') or 'No executive summary was produced.'}",
                    f"Decision supported: {_text(brief.get('decision'), 'The decision owner was not specified.')}",
                    f"Recommended response: {first_action}",
                ])
            ]},
            {"heading": "What was measured", "body": "", "blocks": [measurement_table]},
            {"heading": "Key findings and implications", "body": "", "blocks": [finding_table]},
            {"heading": "Recommended next steps", "body": "", "blocks": [recommendation_table, _bullets("Monitoring plan", result.get("monitoring_metrics"))]},
            {"heading": "Further questions", "body": "", "blocks": [_bullets("Open questions", result.get("unanswered_questions"))]},
            {"heading": "Caveats and assumptions", "body": "", "blocks": [_bullets("Limitations", result.get("limitations"))]},
            {"heading": "Release readiness", "body": "", "blocks": [gate_table]},
        ]
        subtitle = "Editable, evidence-linked decision report for leadership review"

    return {
        "schema_version": "2.0",
        "document_type": document_type,
        "title": session.business_task if document_type in {"executive_report", "case_study_report"} else title,
        "subtitle": subtitle,
        "brand": {"organization": "AI Analytics Workspace", "prepared_for": "Decision owner", "classification": "Review before distribution"},
        "sections": sections,
    }


def _font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    if shading.getparent() is None:
        tc_pr.append(shading)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:cantSplit")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def _set_table_geometry(table, widths: Sequence[int]) -> None:
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    if tbl_w.getparent() is None:
        table_pr.append(tbl_w)
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    if tbl_ind.getparent() is None:
        table_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            if tc_w.getparent() is None:
                tc_pr.append(tc_w)
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
            if margins.getparent() is None:
                tc_pr.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                if node.getparent() is None:
                    margins.append(node)


def _column_widths(count: int) -> list[int]:
    if count <= 1:
        return [CONTENT_WIDTH_DXA]
    if count == 2:
        return [2800, 6560]
    if count == 3:
        return [1800, 2800, 4760]
    if count == 4:
        return [1300, 2500, 2780, 2780]
    first = 1000
    remaining = CONTENT_WIDTH_DXA - first
    base = remaining // (count - 1)
    widths = [first] + [base] * (count - 1)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _add_table(document: Document, block: dict) -> None:
    columns = [str(value) for value in block.get("columns") or []]
    if not columns:
        return
    rows = block.get("rows") or []
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    _set_repeat_table_header(table.rows[0])
    _prevent_row_split(table.rows[0])
    for index, label in enumerate(columns):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, "173461")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].paragraph_format.keep_with_next = True
        run = cell.paragraphs[0].add_run(label)
        _font(run, size=8.5, color="FFFFFF", bold=True)
    for row_index, values in enumerate(rows):
        new_row = table.add_row()
        _prevent_row_split(new_row)
        cells = new_row.cells
        for column_index in range(len(columns)):
            value = values[column_index] if column_index < len(values) else ""
            run = cells[column_index].paragraphs[0].add_run(str(value))
            _font(run, size=8.5, color="14213D")
            cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2:
                _set_cell_shading(cells[column_index], "F4F7FB")
    _set_table_geometry(table, _column_widths(len(columns)))


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _configure_document(document: Document, document_type: str) -> None:
    compact = document_type in {"analysis_journal", "technical_appendix"}
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    language = normal._element.rPr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        normal._element.rPr.append(language)
    language.set(qn("w:val"), "en-US")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25 if compact else 1.10
    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18 if compact else 16, 10 if compact else 8),
        "Heading 2": (13, "2E74B5", 14 if compact else 12, 7 if compact else 6),
        "Heading 3": (12, "1F4D78", 10 if compact else 8, 5 if compact else 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    bullet = document.styles["List Bullet"]
    bullet.paragraph_format.left_indent = Inches(0.375 if compact else 0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.188 if compact else -0.25)
    bullet.paragraph_format.space_after = Pt(4 if compact else 8)
    bullet.paragraph_format.line_spacing = 1.25 if compact else 1.167


def _render_block(document: Document, block: dict) -> None:
    title = str(block.get("title") or "").strip()
    if title:
        document.add_heading(title, level=2)
    kind = block.get("type")
    if kind == "table":
        _add_table(document, block)
    elif kind == "bullets":
        for item in block.get("items") or []:
            document.add_paragraph(str(item), style="List Bullet")
    else:
        for paragraph in str(block.get("text") or "").splitlines() or [""]:
            document.add_paragraph(paragraph)


def build_docx(content: dict, destination: Path, chart_paths: list[Path] | None = None) -> Path:
    """Create a professionally styled, fully editable DOCX from a saved revision."""
    document = Document()
    document_type = content.get("document_type", "executive_report")
    _configure_document(document, document_type)
    document.core_properties.title = content.get("title", "Analysis document")
    document.core_properties.subject = content.get("subtitle", "Evidence-linked analytics deliverable")
    document.core_properties.author = content.get("brand", {}).get("organization", "AI Analytics Workspace")
    document.core_properties.keywords = "analytics, evidence, Google Data Analytics, decision support"
    document.core_properties.comments = "Review all human edits, source permissions, and limitations before external distribution."
    section = document.sections[0]
    header = section.header.paragraphs[0]
    _font(header.add_run(f"AI ANALYTICS WORKSPACE  |  {document_type.replace('_', ' ').upper()}"), size=8.5, color="607089", bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(footer.add_run("Evidence-linked analysis  |  Page "), size=8.5, color="607089")
    _add_page_number(footer)
    _font(footer.add_run("  |  Review before distribution"), size=8.5, color="607089")

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    _font(kicker.add_run(content.get("brand", {}).get("organization", "AI ANALYTICS WORKSPACE").upper()), size=9, color="087F8C", bold=True)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _font(title.add_run(content["title"]), size=28 if document_type == "executive_report" else 24, color="173461", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    _font(subtitle.add_run(content.get("subtitle") or ""), size=13, color="607089")
    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(16)
    classification = content.get("brand", {}).get("classification", "Review before distribution")
    _font(meta.add_run(f"{classification}  |  Exported {datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M UTC')}"), size=9, color="607089")

    for item in content.get("sections", []):
        heading = item.get("heading") or "Untitled section"
        document.add_heading(heading, level=1)
        body = str(item.get("body") or "").strip()
        if body:
            for line in body.splitlines():
                document.add_paragraph(line)
        for block in item.get("blocks") or []:
            _render_block(document, block)

    existing_charts = [path for path in chart_paths or [] if path.is_file()]
    if existing_charts:
        document.add_heading("Evidence charts", level=1)
        for index, path in enumerate(existing_charts, start=1):
            picture = document.add_picture(str(path), width=Inches(6.2))
            alt_text = f"Evidence chart {index}: {path.stem.replace('_', ' ')}"
            picture._inline.docPr.set("descr", alt_text)
            picture._inline.docPr.set("title", alt_text)
            caption = document.add_paragraph(f"Figure {index}. Evidence chart generated from the validated analysis.")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                _font(run, size=9, color="607089")

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return destination


def build_pdf(content: dict, destination: Path, chart_paths: list[Path] | None = None) -> Path:
    """Create a print-ready PDF from the same saved structured revision."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image, ListFlowable, ListItem, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    destination.parent.mkdir(parents=True, exist_ok=True)
    document = SimpleDocTemplate(str(destination), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=0.82 * inch, bottomMargin=0.75 * inch, title=content.get("title", "Analysis document"), author="AI Analytics Workspace")
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="DocTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=colors.HexColor("#173461"), spaceAfter=7))
    styles.add(ParagraphStyle(name="Subtitle", parent=styles["Normal"], fontSize=11, leading=15, textColor=colors.HexColor("#607089"), spaceAfter=14))
    styles.add(ParagraphStyle(name="H1Brand", parent=styles["Heading1"], fontSize=15, leading=18, textColor=colors.HexColor("#2E74B5"), spaceBefore=14, spaceAfter=7))
    styles.add(ParagraphStyle(name="H2Brand", parent=styles["Heading2"], fontSize=11.5, leading=14, textColor=colors.HexColor("#1F4D78"), spaceBefore=9, spaceAfter=5))
    styles.add(ParagraphStyle(name="BodyBrand", parent=styles["BodyText"], fontSize=9.5, leading=13, textColor=colors.HexColor("#14213D"), spaceAfter=6))
    styles.add(ParagraphStyle(name="TableHeader", parent=styles["BodyText"], fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=colors.white, spaceAfter=0))
    def safe(value: Any) -> str:
        return escape(str(value if value is not None else "").replace("→", " to ")).replace("\n", "<br/>")

    story = [Paragraph(safe(content.get("brand", {}).get("organization", "AI ANALYTICS WORKSPACE").upper()), ParagraphStyle(name="Kicker", parent=styles["Normal"], fontSize=8, textColor=colors.HexColor("#087F8C"), spaceAfter=4)), Paragraph(safe(content.get("title", "Analysis document")), styles["DocTitle"]), Paragraph(safe(content.get("subtitle", "")), styles["Subtitle"]), Paragraph(safe(f"{content.get('brand', {}).get('classification', 'Review before distribution')} · Exported {datetime.now(timezone.utc).strftime('%d %b %Y, %H:%M UTC')}"), styles["BodyBrand"]), Spacer(1, 8)]

    for section_index, section_data in enumerate(content.get("sections") or []):
        if section_index and section_data.get("phase") in {"Prepare", "Process", "Analyze", "Share", "Act"}:
            story.append(PageBreak())
        story.append(Paragraph(safe(section_data.get("heading") or "Untitled section"), styles["H1Brand"]))
        if section_data.get("body"):
            story.append(Paragraph(safe(section_data["body"]), styles["BodyBrand"]))
        for block in section_data.get("blocks") or []:
            if block.get("title"):
                story.append(Paragraph(safe(block["title"]), styles["H2Brand"]))
            if block.get("type") == "table" and block.get("columns"):
                rows = [[Paragraph(safe(value), styles["TableHeader"]) for value in block["columns"]]]
                rows.extend([[Paragraph(safe(value), styles["BodyBrand"]) for value in row] for row in (block.get("rows") or [])])
                count = len(block["columns"])
                col_width = document.width / max(count, 1)
                table = Table(rows, colWidths=[col_width] * count, repeatRows=1, hAlign="LEFT")
                table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#173461")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D9E1EC")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F4F7FB")])]))
                story.extend([table, Spacer(1, 7)])
            elif block.get("type") == "bullets":
                story.append(ListFlowable([ListItem(Paragraph(safe(value), styles["BodyBrand"])) for value in block.get("items") or []], bulletType="bullet", leftIndent=18))
            else:
                story.append(Paragraph(safe(block.get("text") or ""), styles["BodyBrand"]))

    for index, path in enumerate([item for item in chart_paths or [] if item.is_file()], start=1):
        story.extend([Paragraph("Evidence charts" if index == 1 else "", styles["H1Brand"]), Image(str(path), width=6.2 * inch, height=3.5 * inch), Paragraph(f"Figure {index}. Evidence chart generated from the validated analysis.", ParagraphStyle(name=f"Caption{index}", parent=styles["BodyBrand"], alignment=TA_CENTER, textColor=colors.HexColor("#607089")))])

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#607089"))
        canvas.drawCentredString(letter[0] / 2, 0.38 * inch, f"Evidence-linked analysis · Page {doc.page} · Review before distribution")
        canvas.restoreState()

    document.build(story, onFirstPage=footer, onLaterPages=footer)
    return destination
